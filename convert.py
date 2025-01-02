from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
import os

# Create a new Word document
doc = Document()

# Read the HTML file
with open('temp_form.html', 'r', encoding='utf-8') as file:
    html_content = file.read()

# Parse HTML
soup = BeautifulSoup(html_content, 'html.parser')

# Function to add text with proper formatting
def add_text(element, doc):
    if element.name in ['h1', 'h2', 'h3']:
        p = doc.add_paragraph()
        run = p.add_run(element.get_text())
        run.bold = True
        font_size = {'h1': 24, 'h2': 20, 'h3': 16}
        run.font.size = Pt(font_size[element.name])
    elif element.name == 'p':
        p = doc.add_paragraph()
        # Handle strong tags within paragraphs
        for content in element.contents:
            if hasattr(content, 'name') and content.name == 'strong':
                run = p.add_run(content.get_text())
                run.bold = True
            elif hasattr(content, 'name') and content.name == 'br':
                run = p.add_run('\n')
            else:
                run = p.add_run(str(content))
    elif element.name == 'ul' or element.name == 'ol':
        for li in element.find_all('li'):
            p = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
            p.add_run(li.get_text())
    elif element.name == 'table':
        table = doc.add_table(rows=1, cols=len(element.find('tr').find_all(['th', 'td'])))
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        for i, th in enumerate(element.find('tr').find_all(['th', 'td'])):
            header_cells[i].text = th.get_text()
            
        # Add data rows
        for tr in element.find_all('tr')[1:]:
            row_cells = table.add_row().cells
            for i, td in enumerate(tr.find_all('td')):
                row_cells[i].text = td.get_text()

# Process each element
for element in soup.body.children:
    if hasattr(element, 'name') and element.name:  # Skip NavigableString objects
        add_text(element, doc)

# Save the document
doc.save('filled_form.docx')

# Clean up temporary HTML file
os.remove('temp_form.html')
